from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = "docs/Guia_explicacion_diapositivas_CowTrack_MVP.docx"


GREEN = "23452C"
LIGHT_GREEN = "E7EED8"
CREAM = "FFF9EF"
STRAW = "D9B85F"
BROWN = "7A4A2D"
DARK = "172719"


slides = [
    {
        "n": "1",
        "title": "Carátula - CowTrack MVP",
        "shows": "Presenta el nombre del proyecto, el alcance general y los autores.",
        "say": (
            "CowTrack es un MVP de visión computacional aplicado a ganadería. "
            "El sistema procesa video aéreo para detectar vacas, trackearlas, contarlas "
            "y reidentificar tres animales catalogados: Marta, Maria y Margarita."
        ),
        "key": "El proyecto está completo y el resultado final es un video HD auditable.",
    },
    {
        "n": "2",
        "title": "Problema y oportunidad",
        "shows": "Explica el contexto agropecuario y las dificultades del seguimiento manual.",
        "say": (
            "En el campo, el conteo y seguimiento individual suelen hacerse de forma manual. "
            "El video con dron aporta mucha información, pero también introduce oclusiones, "
            "giros de cámara, cambios de escala y animales visualmente parecidos."
        ),
        "key": "La oportunidad es convertir video aéreo en información operativa.",
    },
    {
        "n": "3",
        "title": "Objetivo del MVP",
        "shows": "Define qué debe lograr el sistema y separa conteo general de reidentificación.",
        "say": (
            "El objetivo fue generar un render HD con bounding boxes para las vacas detectadas, "
            "etiquetas priorizadas para las tres vacas catalogadas y métricas finales. "
            "La reidentificación individual y el conteo general se evalúan como problemas relacionados, pero distintos."
        ),
        "key": "El núcleo del MVP es reconocer y sostener identidades individuales en el tiempo.",
    },
    {
        "n": "4",
        "title": "Datos, modelos y galería",
        "shows": "Resume el uso de OpenCows, el modelo Re-ID general y la galería Erondina.",
        "say": (
            "Primero se entrenó un extractor Re-ID general con OpenCows. "
            "Luego se construyó una galería específica con fotos extraídas del campo Erondina, "
            "aislando mejor las referencias de Marta, Maria y Margarita para comparar embeddings."
        ),
        "key": "El modelo general genera embeddings; la galería Erondina permite decidir identidades puntuales.",
    },
    {
        "n": "5",
        "title": "Arquitectura del pipeline final",
        "shows": "Muestra el flujo completo: video, YOLOv8, tracking, recortes, embeddings, FAISS y render.",
        "say": (
            "YOLOv8 detecta las vacas por frame. El tracker genera trayectorias base. "
            "Luego se recortan las vacas y se extraen embeddings de cuerpo y subembeddings regionales, "
            "como cabeza o zona superior. FAISS compara esos vectores contra la galería Erondina y permite "
            "bloquear identidades para el render final."
        ),
        "key": "El render final no decide todo desde cero en cada frame: sigue temporalmente la identidad reconocida.",
    },
    {
        "n": "6",
        "title": "Validación Re-ID: tres identidades",
        "shows": "Presenta la evidencia visual y los scores finales de Marta, Maria y Margarita.",
        "say": (
            "En la validación final, las tres identidades objetivo fueron reconocidas correctamente. "
            "Los scores finales fueron altos: Margarita 0.9269, Maria 0.8973 y Marta 0.9459. "
            "Para esta validación de identidad, la precisión, el recall y el mAP@0.5 Re-ID quedaron en 100%."
        ),
        "key": "La reidentificación individual cumple el criterio principal del MVP.",
    },
    {
        "n": "7",
        "title": "Resultado visual del render HD",
        "shows": "Incluye capturas del video final con etiquetas y bounding boxes.",
        "say": (
            "Esta diapositiva muestra el resultado observable: las vacas catalogadas aparecen con etiquetas destacadas "
            "y las demás vacas se muestran con cajas de detección. La prioridad visual está puesta en Marta, Maria y Margarita."
        ),
        "key": "El resultado final es interpretable y sirve como evidencia visual del funcionamiento.",
    },
    {
        "n": "8",
        "title": "Métricas del MVP final",
        "shows": "Resume duración, resolución, presencia temporal, estabilidad y métricas Re-ID.",
        "say": (
            "El video final tiene 47.01 segundos, 1409 frames, 29.97 FPS y resolución 1920x1080. "
            "Las identidades catalogadas tuvieron una presencia promedio de 99.81% y no se registraron ID switches "
            "en la validación final."
        ),
        "key": "Las métricas fuertes del MVP están en la reidentificación y continuidad de las vacas catalogadas.",
    },
    {
        "n": "9",
        "title": "Conteo general: lectura honesta",
        "shows": "Explica la diferencia entre 13 vacas reales y 21 etiquetas automáticas.",
        "say": (
            "El fragmento final tiene 13 vacas reales confirmadas visualmente, pero el tracker genera 21 etiquetas automáticas. "
            "Esto no invalida la Re-ID de las tres vacas catalogadas; muestra una limitación del conteo general, "
            "donde algunas vacas no catalogadas se fragmentan en más de una etiqueta."
        ),
        "key": "La Re-ID individual se valida con éxito; el conteo automático general queda como mejora futura.",
    },
    {
        "n": "10",
        "title": "Conclusiones y próximos pasos",
        "shows": "Cierra el proyecto, declara el MVP completo y propone mejoras.",
        "say": (
            "CowTrack demuestra que es posible aplicar IA y visión computacional a un problema real de monitoreo ganadero. "
            "El MVP queda completo porque integra detección, tracking, Re-ID, métricas y render final. "
            "Como próximos pasos, se propone mejorar el conteo de vacas no catalogadas, ampliar la galería y anotar ground truth "
            "para medir métricas de detección con mayor precisión."
        ),
        "key": "El MVP resuelve la reidentificación principal y deja una línea clara de evolución.",
    },
]


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=DARK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Aptos Display"
    run.font.color.rgb = RGBColor.from_string(GREEN if level == 1 else BROWN)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.6)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Guía breve para explicar las 10 diapositivas")
run.bold = True
run.font.name = "Aptos Display"
run.font.size = Pt(24)
run.font.color.rgb = RGBColor.from_string(GREEN)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("CowTrack MVP - Detección, conteo, tracking y reidentificación de vacas")
run.font.name = "Aptos"
run.font.size = Pt(12)
run.font.color.rgb = RGBColor.from_string(BROWN)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Instituto de Formación Técnica Superior Número 11 · Ciencia de Datos e Inteligencia Artificial")
run.font.name = "Aptos"
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor.from_string("666666")

intro = doc.add_paragraph()
intro.paragraph_format.space_before = Pt(14)
intro.paragraph_format.space_after = Pt(10)
intro.add_run(
    "Este documento funciona como guion de apoyo para presentar la PPT del proyecto CowTrack MVP. "
    "Cada diapositiva incluye una explicación breve, el mensaje técnico que conviene destacar y la idea central "
    "que debería quedar clara para la audiencia."
)

add_heading(doc, "Resumen de exposición", 1)
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
headers = ["Slide", "Título", "Qué explica", "Mensaje clave"]
for i, header in enumerate(headers):
    shade_cell(table.rows[0].cells[i], GREEN)
    set_cell_text(table.rows[0].cells[i], header, bold=True, color="FFFFFF", size=9)

for item in slides:
    row = table.add_row().cells
    set_cell_text(row[0], item["n"], bold=True, color=GREEN, size=10)
    set_cell_text(row[1], item["title"], bold=True, color=DARK, size=9)
    set_cell_text(row[2], item["shows"], color=DARK, size=8.5)
    set_cell_text(row[3], item["key"], color=DARK, size=8.5)
    for cell in row:
        shade_cell(cell, CREAM)

add_heading(doc, "Guion breve por diapositiva", 1)

for item in slides:
    add_heading(doc, f"Diapositiva {item['n']} - {item['title']}", 2)
    p1 = doc.add_paragraph()
    p1.add_run("Qué muestra: ").bold = True
    p1.add_run(item["shows"])
    p2 = doc.add_paragraph()
    p2.add_run("Cómo explicarla: ").bold = True
    p2.add_run(item["say"])
    p3 = doc.add_paragraph()
    p3.add_run("Idea clave: ").bold = True
    p3.add_run(item["key"])
    p3.paragraph_format.space_after = Pt(8)

add_heading(doc, "Cierre sugerido", 1)
closing = doc.add_paragraph()
closing.add_run(
    "CowTrack MVP integra detección con YOLOv8, tracking, embeddings Re-ID, comparación vectorial con FAISS "
    "y render HD final. La principal fortaleza del proyecto es la reidentificación estable de Marta, Maria y Margarita. "
    "El conteo general queda documentado como resultado operativo aceptable y como línea futura de mejora."
)

doc.save(OUTPUT)
print(OUTPUT)
